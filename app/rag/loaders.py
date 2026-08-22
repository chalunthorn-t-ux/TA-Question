"""อ่านเอกสารหลายรูปแบบให้กลายเป็น "ชิ้นข้อความ + metadata" รูปแบบเดียวกัน

รองรับ: .pdf .docx .xlsx .xls .csv .txt .md
แต่ละ loader คืน list[RawSection] โดย section = ข้อความ 1 ก้อนที่ยังไม่ถูกตัด chunk
"""

from __future__ import annotations

import csv
import re
from dataclasses import dataclass, field
from pathlib import Path
from typing import Iterable


@dataclass
class RawSection:
    """ข้อความก้อนหนึ่งจากเอกสาร ก่อนนำไปตัด chunk"""

    text: str
    source: str                                  # ชื่อไฟล์
    locator: str = ""                            # หน้า / ชีต / แถว — ใช้อ้างอิงตอนแสดงผล
    meta: dict = field(default_factory=dict)


class UnsupportedFileError(Exception):
    pass


def clean_text(text: str) -> str:
    """ลดช่องว่างซ้ำซ้อน แต่ยังคงย่อหน้าไว้"""
    text = text.replace(" ", " ").replace("\r\n", "\n").replace("\r", "\n")
    text = re.sub(r"[ \t]+", " ", text)
    text = re.sub(r"\n{3,}", "\n\n", text)
    return text.strip()


# --------------------------------------------------------------------------- #
# PDF
# --------------------------------------------------------------------------- #
def load_pdf(path: Path) -> list[RawSection]:
    from pypdf import PdfReader

    sections: list[RawSection] = []
    reader = PdfReader(str(path))
    for page_no, page in enumerate(reader.pages, start=1):
        try:
            text = clean_text(page.extract_text() or "")
        except Exception:
            text = ""
        if text:
            sections.append(
                RawSection(text=text, source=path.name, locator=f"หน้า {page_no}")
            )
    return sections


# --------------------------------------------------------------------------- #
# Word
# --------------------------------------------------------------------------- #
def load_docx(path: Path) -> list[RawSection]:
    import docx

    doc = docx.Document(str(path))
    sections: list[RawSection] = []

    # ย่อหน้าปกติ — จัดกลุ่มตามหัวข้อ (Heading) เพื่อให้ context ไม่หลุดจากกัน
    buffer: list[str] = []
    heading = ""

    def flush() -> None:
        if buffer:
            body = clean_text("\n".join(buffer))
            if body:
                sections.append(
                    RawSection(
                        text=(f"{heading}\n{body}" if heading else body),
                        source=path.name,
                        locator=heading or "เนื้อหา",
                    )
                )
            buffer.clear()

    for para in doc.paragraphs:
        txt = (para.text or "").strip()
        if not txt:
            continue
        if para.style is not None and str(para.style.name).startswith("Heading"):
            flush()
            heading = txt
        else:
            buffer.append(txt)
    flush()

    # ตารางใน Word — แปลงเป็นข้อความแบบ key: value เพื่อให้ค้นเจอ
    for t_idx, table in enumerate(doc.tables, start=1):
        rows = [[cell.text.strip() for cell in row.cells] for row in table.rows]
        rendered = _render_table(rows)
        if rendered:
            sections.append(
                RawSection(
                    text=rendered, source=path.name, locator=f"ตารางที่ {t_idx}"
                )
            )
    return sections


def _render_table(rows: list[list[str]]) -> str:
    """แปลงตารางเป็นข้อความอ่านง่าย: ใช้แถวแรกเป็นหัวคอลัมน์"""
    rows = [r for r in rows if any(c for c in r)]
    if not rows:
        return ""
    header, *body = rows
    if not body:
        return " | ".join(header)
    lines: list[str] = []
    for row in body:
        pairs = [
            f"{(header[i] if i < len(header) else f'คอลัมน์{i + 1}')}: {val}"
            for i, val in enumerate(row)
            if val
        ]
        if pairs:
            lines.append(" | ".join(pairs))
    return "\n".join(lines)


# --------------------------------------------------------------------------- #
# Excel / CSV
# --------------------------------------------------------------------------- #
_QUESTION_COLS = {"question", "คำถาม", "หัวข้อ", "q", "ประเด็น"}
_ANSWER_COLS = {"answer", "คำตอบ", "รายละเอียด", "a", "คำอธิบาย"}


def _read_csv_sheet(path: Path) -> list[list[str]]:
    encoding = _sniff_encoding(path)
    with path.open("r", encoding=encoding, newline="") as fh:
        sample = fh.read(8192)
        fh.seek(0)
        try:
            dialect = csv.Sniffer().sniff(sample, delimiters=",;\t|")
        except csv.Error:
            dialect = csv.excel
        return [[(cell or "").strip() for cell in row] for row in csv.reader(fh, dialect)]


def _read_excel_sheets(path: Path) -> dict[str, list[list[str]]]:
    from openpyxl import load_workbook

    book = load_workbook(str(path), read_only=True, data_only=True)
    try:
        sheets: dict[str, list[list[str]]] = {}
        for ws in book.worksheets:
            rows = [
                ["" if cell is None else str(cell).strip() for cell in row]
                for row in ws.iter_rows(values_only=True)
            ]
            sheets[str(ws.title)] = rows
        return sheets
    finally:
        book.close()


def _read_xls_sheets(path: Path) -> dict[str, list[list[str]]]:
    """Excel รุ่นเก่า (.xls) — openpyxl อ่านไม่ได้ ต้องใช้ xlrd"""
    try:
        import xlrd
    except ImportError as exc:
        raise UnsupportedFileError(
            "ต้องติดตั้ง xlrd เพื่ออ่านไฟล์ .xls (หรือบันทึกไฟล์เป็น .xlsx แทน)"
        ) from exc

    book = xlrd.open_workbook(str(path))
    return {
        sheet.name: [
            ["" if v is None else str(v).strip() for v in sheet.row_values(r)]
            for r in range(sheet.nrows)
        ]
        for sheet in book.sheets()
    }


def load_tabular(path: Path) -> list[RawSection]:
    """Excel/CSV — ถ้าเป็นตาราง FAQ (มีคอลัมน์คำถาม/คำตอบ) จะจับคู่ให้เป็น 1 chunk ต่อ 1 แถว"""
    suffix = path.suffix.lower()
    if suffix == ".csv":
        sheets = {"CSV": _read_csv_sheet(path)}
    elif suffix == ".xls":
        sheets = _read_xls_sheets(path)
    else:
        sheets = _read_excel_sheets(path)

    sections: list[RawSection] = []
    for sheet_name, rows in sheets.items():
        rows = [r for r in rows if any(c for c in r)]
        if len(rows) < 2:                      # มีแต่หัวตาราง หรือว่างเปล่า
            continue

        header = [c.strip() or f"คอลัมน์{i + 1}" for i, c in enumerate(rows[0])]
        lookup = {h.lower(): i for i, h in enumerate(header)}
        q_idx = next((lookup[k] for k in lookup if k in _QUESTION_COLS), None)
        a_idx = next((lookup[k] for k in lookup if k in _ANSWER_COLS), None)

        for row_no, row in enumerate(rows[1:], start=2):   # แถว 1 = header
            cells = [c.strip() for c in row] + [""] * (len(header) - len(row))
            locator = f"{sheet_name} · แถว {row_no}"

            if q_idx is not None and a_idx is not None:
                question = cells[q_idx] if q_idx < len(cells) else ""
                answer = cells[a_idx] if a_idx < len(cells) else ""
                if not (question or answer):
                    continue
                text = f"คำถาม: {question}\nคำตอบ: {answer}"
                extra = [
                    f"{header[i]}: {v}"
                    for i, v in enumerate(cells)
                    if i not in (q_idx, a_idx) and i < len(header) and v
                ]
                if extra:
                    text += "\n" + " | ".join(extra)
            else:
                pairs = [
                    f"{header[i]}: {v}"
                    for i, v in enumerate(cells)
                    if i < len(header) and v
                ]
                if not pairs:
                    continue
                text = " | ".join(pairs)

            sections.append(
                RawSection(text=clean_text(text), source=path.name, locator=locator)
            )
    return sections


def _sniff_encoding(path: Path) -> str:
    """CSV ไทยมักเป็น utf-8-sig หรือ cp874 — ลองไล่ดู"""
    for enc in ("utf-8-sig", "utf-8", "cp874"):
        try:
            with path.open("r", encoding=enc) as fh:
                csv.reader(fh).__next__()
            return enc
        except (UnicodeDecodeError, StopIteration):
            continue
    return "utf-8"


# --------------------------------------------------------------------------- #
# ข้อความล้วน
# --------------------------------------------------------------------------- #
def load_text(path: Path) -> list[RawSection]:
    for enc in ("utf-8-sig", "utf-8", "cp874"):
        try:
            raw = path.read_text(encoding=enc)
            break
        except UnicodeDecodeError:
            continue
    else:
        return []

    # ตัดตามหัวข้อ markdown ถ้ามี
    blocks = re.split(r"\n(?=#{1,3} )", raw)
    sections: list[RawSection] = []
    for block in blocks:
        text = clean_text(block)
        if not text:
            continue
        first_line = text.splitlines()[0].lstrip("# ").strip()
        sections.append(
            RawSection(text=text, source=path.name, locator=first_line[:60] or "เนื้อหา")
        )
    return sections


# --------------------------------------------------------------------------- #
LOADERS = {
    ".pdf": load_pdf,
    ".docx": load_docx,
    ".xlsx": load_tabular,
    ".xls": load_tabular,
    ".csv": load_tabular,
    ".txt": load_text,
    ".md": load_text,
}


def load_file(path: Path) -> list[RawSection]:
    loader = LOADERS.get(path.suffix.lower())
    if loader is None:
        raise UnsupportedFileError(f"ยังไม่รองรับไฟล์ประเภท {path.suffix}")
    return loader(path)


def iter_documents(root: Path) -> Iterable[Path]:
    """ไล่ไฟล์ที่รองรับทั้งหมดในโฟลเดอร์ (รวมโฟลเดอร์ย่อย)"""
    for path in sorted(root.rglob("*")):
        if path.is_file() and path.suffix.lower() in LOADERS and not path.name.startswith("~$"):
            yield path
